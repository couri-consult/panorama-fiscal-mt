"""Generate an executive Word report (.docx) from data.json.

Output: docs/relatorio_executivo.docx
Uses python-docx (pip install python-docx).

Run:
    python scripts/build_report.py
"""
import json
import os
import datetime

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.dirname(HERE)
DATA = os.path.join(ROOT, "data.json")
OUTPUT = os.path.join(ROOT, "docs", "agenda-politica-fiscal.docx")

# Theme — alinhado com o painel (verde STN + neutros)
COLOR_HEADER_DARK = RGBColor(0x1B, 0x5E, 0x20)   # --g800
COLOR_HEADER_MID = RGBColor(0x2E, 0x7D, 0x32)    # --g700
COLOR_ACCENT = RGBColor(0x4C, 0xAF, 0x50)        # --g500
COLOR_RED = RGBColor(0xC0, 0x39, 0x2B)
COLOR_AMBER = RGBColor(0xC4, 0x7A, 0x10)
COLOR_BLUE = RGBColor(0x1A, 0x5F, 0xA6)
COLOR_MUTED = RGBColor(0x6B, 0x72, 0x80)
COLOR_TEXT = RGBColor(0x1A, 0x1A, 0x1A)


def cite_rreo(period_str):
    """'2026/1' -> 'RREO – 1º bimestre de 2026'."""
    year, bim = period_str.split("/")
    return f"RREO – {bim}º bimestre de {year}"


def cite_rgf(period_str, poder="Poder Executivo"):
    """'2025/3' -> 'RGF – Poder Executivo – 3º quadrimestre de 2025'."""
    year, quad = period_str.split("/")
    return f"RGF – {poder} – {quad}º quadrimestre de {year}"


def fmt_bi(reais, decimals=1):
    return f"R$ {reais/1e9:.{decimals}f} bi".replace(".", ",")


def fmt_mi(reais):
    return f"R$ {reais/1e6:.0f} mi"


def fmt_pct(x, decimals=1):
    return f"{x:.{decimals}f}%".replace(".", ",")


def shade_cell(cell, color_hex):
    """Adiciona cor de fundo a uma célula (hex sem #)."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def set_cell_borders(cell, color="CCCCCC", size=4):
    """Define bordas finas em toda volta da célula."""
    tc_pr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), str(size))
        b.set(qn("w:color"), color)
        tcBorders.append(b)
    tc_pr.append(tcBorders)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        if level == 0:
            run.font.color.rgb = COLOR_HEADER_DARK
            run.font.size = Pt(24)
        elif level == 1:
            run.font.color.rgb = COLOR_HEADER_MID
            run.font.size = Pt(16)
        elif level == 2:
            run.font.color.rgb = COLOR_HEADER_MID
            run.font.size = Pt(13)
    return h


def add_paragraph(doc, text, bold=False, italic=False, size=11, color=None, align=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    if color:
        r.font.color.rgb = color
    if align:
        p.alignment = align
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    r.font.size = Pt(11)
    return p


def add_source(doc, text):
    """Linha pequena em itálico cinza para 'Fonte: ...' abaixo de tabela ou parágrafo."""
    p = doc.add_paragraph()
    r = p.add_run(f"Fonte: {text}")
    r.font.size = Pt(9)
    r.font.italic = True
    r.font.color.rgb = COLOR_MUTED
    return p


def add_kv_paragraph(doc, label, value, value_color=None):
    """Paragraph with bold label and plain value."""
    p = doc.add_paragraph()
    rl = p.add_run(f"{label} ")
    rl.font.bold = True
    rl.font.size = Pt(11)
    rv = p.add_run(value)
    rv.font.size = Pt(11)
    if value_color:
        rv.font.color.rgb = value_color
    return p


def add_kpi_table(doc, kpis):
    """Tabela 2 colunas: indicador → valor + sub."""
    labels = {
        "orcamento_df": "Orçamento DF (próprio)",
        "fcdf": "FCDF (Fundo Constitucional)",
        "investimento_rcl": "Investimento / RCL",
        "capag": "Nota CAPAG consolidada",
        "caixa": "Disponibilidade de caixa líquida",
        "beneficios": "Benefícios tributários",
    }
    color_map = {"red": COLOR_RED, "amber": COLOR_AMBER, "blue": COLOR_BLUE, "green": COLOR_HEADER_MID}

    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    # Header row
    hdr = table.rows[0].cells
    hdr[0].text = ""
    hdr[1].text = ""
    rh0 = hdr[0].paragraphs[0].add_run("Indicador")
    rh1 = hdr[1].paragraphs[0].add_run("Valor")
    for r in (rh0, rh1):
        r.font.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.size = Pt(11)
    for c in hdr:
        shade_cell(c, "1B5E20")
        set_cell_borders(c, "1B5E20")

    for kpi in kpis:
        ch = kpi.get("chave")
        label = labels.get(ch, ch)
        row = table.add_row().cells
        row[0].text = ""
        row[1].text = ""
        # label column
        p0 = row[0].paragraphs[0]
        r_lbl = p0.add_run(label)
        r_lbl.font.size = Pt(11)
        r_lbl.font.bold = True
        # value column: big value + sub small below
        p1 = row[1].paragraphs[0]
        r_val = p1.add_run(str(kpi.get("valor_bilhoes") or ""))
        r_val.font.size = Pt(13)
        r_val.font.bold = True
        cor = color_map.get(kpi.get("cor"), COLOR_TEXT)
        r_val.font.color.rgb = cor
        sub = kpi.get("sub")
        if sub:
            p_sub = row[1].add_paragraph()
            r_sub = p_sub.add_run(sub)
            r_sub.font.size = Pt(9)
            r_sub.font.italic = True
            r_sub.font.color.rgb = COLOR_MUTED
        for c in row:
            set_cell_borders(c)
    # column widths
    for row in table.rows:
        row.cells[0].width = Cm(7)
        row.cells[1].width = Cm(9)


def add_capag_table(doc, capag_rows, capag_hist, nota_consolidada):
    """Tabela CAPAG indicador × nota × valor × situação."""
    table = doc.add_table(rows=1, cols=4)
    table.autofit = False
    hdr = table.rows[0].cells
    headers = ["Indicador", "Nota", "Valor", "Situação"]
    for c, txt in zip(hdr, headers):
        c.text = ""
        r = c.paragraphs[0].add_run(txt)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.size = Pt(10)
        shade_cell(c, "2E7D32")
        set_cell_borders(c, "2E7D32")

    badge_color = {"green": COLOR_HEADER_MID, "amber": COLOR_AMBER, "red": COLOR_RED}
    for row in capag_rows:
        cells = table.add_row().cells
        cells[0].paragraphs[0].add_run(row.get("indicador", "")).font.size = Pt(10)
        # nota colorida
        nota_run = cells[1].paragraphs[0].add_run(row.get("nota", ""))
        nota_run.font.bold = True
        nota_run.font.size = Pt(12)
        nota_run.font.color.rgb = badge_color.get(row.get("badge", ""), COLOR_TEXT)
        cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        # valor
        cells[2].paragraphs[0].add_run(row.get("valor", "")).font.size = Pt(10)
        cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        # situação
        sit_run = cells[3].paragraphs[0].add_run(row.get("obs", ""))
        sit_run.font.size = Pt(9)
        sit_run.font.italic = True
        sit_run.font.color.rgb = COLOR_MUTED
        for c in cells:
            set_cell_borders(c)

    # column widths
    widths = [Cm(6), Cm(2), Cm(3), Cm(6)]
    for row in table.rows:
        for cell, w in zip(row.cells, widths):
            cell.width = w


def add_caixa_table(doc, caixa_rows):
    table = doc.add_table(rows=1, cols=3)
    table.autofit = False
    hdr = table.rows[0].cells
    for c, txt in zip(hdr, ["Ano", "Caixa líquido (após RPNP)", "Rec. não vinculados"]):
        c.text = ""
        r = c.paragraphs[0].add_run(txt)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.size = Pt(10)
        shade_cell(c, "2E7D32")
        set_cell_borders(c, "2E7D32")

    for row in caixa_rows:
        cells = table.add_row().cells
        cells[0].paragraphs[0].add_run(str(row.get("ano", ""))).font.size = Pt(10)
        cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        liq = row.get("caixa_liquido", 0) or 0
        nv = row.get("nao_vinculado", 0) or 0
        for i, val in enumerate([liq, nv], 1):
            r = cells[i].paragraphs[0].add_run(fmt_bi(val, decimals=2) if abs(val) >= 1e9 else fmt_mi(val))
            r.font.size = Pt(10)
            cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            if val < 0:
                r.font.color.rgb = COLOR_RED
        for c in cells:
            set_cell_borders(c)

    for row in table.rows:
        for cell, w in zip(row.cells, [Cm(2), Cm(7), Cm(7)]):
            cell.width = w


def add_capag_hist_table(doc, hist):
    table = doc.add_table(rows=1, cols=5)
    table.autofit = False
    hdr_text = ["Ano", "Endividamento", "Poup. Corrente", "Liquidez", "Consolidado"]
    for c, txt in zip(table.rows[0].cells, hdr_text):
        c.text = ""
        r = c.paragraphs[0].add_run(txt)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.size = Pt(9)
        shade_cell(c, "2E7D32")
        set_cell_borders(c, "2E7D32")
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    color_map = {"A": COLOR_HEADER_MID, "B": COLOR_AMBER, "C": COLOR_RED, "D": COLOR_RED, "ND": COLOR_MUTED}
    for row in hist:
        cells = table.add_row().cells
        vals = [row["ano"], row["endividamento"], row["poupanca"], row["liquidez"], row["consolidado"]]
        for i, v in enumerate(vals):
            cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = cells[i].paragraphs[0].add_run(str(v))
            r.font.size = Pt(10)
            if i > 0:
                r.font.bold = True
                r.font.color.rgb = color_map.get(v, COLOR_TEXT)
        for c in cells:
            set_cell_borders(c)
    widths = [Cm(2), Cm(3.5), Cm(3.5), Cm(3), Cm(3.5)]
    for row in table.rows:
        for cell, w in zip(row.cells, widths):
            cell.width = w


def add_pessoal_table(doc, pessoal_rows):
    pmap = {r["chave"]: r["valor"] for r in pessoal_rows}
    atual = pmap.get("atual_pct", 0)
    alerta = pmap.get("alerta_pct", 44.1)
    prudencial = pmap.get("prudencial_pct", 46.55)
    maximo = pmap.get("maximo_pct", 49)
    rcl = pmap.get("rcl_bi", 39.2)

    table = doc.add_table(rows=1, cols=3)
    table.autofit = False
    for c, txt in zip(table.rows[0].cells, ["Limite", "% da RCL", "Espaço fiscal"]):
        c.text = ""
        r = c.paragraphs[0].add_run(txt)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.size = Pt(10)
        shade_cell(c, "2E7D32")
        set_cell_borders(c, "2E7D32")

    def make_row(label, pct, space_bi, color=None, bold_left=False):
        cells = table.add_row().cells
        cells[0].text = ""
        rl = cells[0].paragraphs[0].add_run(label)
        rl.font.size = Pt(10)
        rl.font.bold = bold_left

        cells[1].text = ""
        rp = cells[1].paragraphs[0].add_run(fmt_pct(pct))
        rp.font.size = Pt(10)
        rp.font.bold = bold_left
        if color: rp.font.color.rgb = color
        cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        cells[2].text = ""
        rs = cells[2].paragraphs[0].add_run(space_bi)
        rs.font.size = Pt(10)
        cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for c in cells:
            set_cell_borders(c)

    def esp(limit):
        bi = (limit - atual) / 100 * rcl
        return f"R$ {bi:.1f} bi".replace(".", ",")

    make_row("Atual", atual, "—", color=COLOR_HEADER_MID, bold_left=True)
    make_row("Alerta", alerta, esp(alerta), color=COLOR_AMBER)
    make_row("Prudencial", prudencial, esp(prudencial), color=COLOR_AMBER)
    make_row("Máximo", maximo, esp(maximo), color=COLOR_RED)

    for row in table.rows:
        for cell, w in zip(row.cells, [Cm(4), Cm(3), Cm(7)]):
            cell.width = w


def add_ppp_table(doc, ppp_rows):
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    for c, txt in zip(table.rows[0].cells, ["Projeto", "Status"]):
        c.text = ""
        r = c.paragraphs[0].add_run(txt)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.size = Pt(10)
        shade_cell(c, "2E7D32")
        set_cell_borders(c, "2E7D32")
    status_color = {"Contratada": COLOR_HEADER_MID, "A contratar": COLOR_BLUE, "Suspenso": COLOR_AMBER}
    for row in ppp_rows:
        cells = table.add_row().cells
        cells[0].text = ""
        cells[0].paragraphs[0].add_run(row.get("nome", "")).font.size = Pt(10)
        cells[1].text = ""
        r = cells[1].paragraphs[0].add_run(row.get("status", ""))
        r.font.size = Pt(10)
        r.font.color.rgb = status_color.get(row.get("status", ""), COLOR_TEXT)
        r.font.bold = True
        cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for c in cells:
            set_cell_borders(c)
    for row in table.rows:
        for cell, w in zip(row.cells, [Cm(11), Cm(5)]):
            cell.width = w


def build_report():
    with open(DATA, encoding="utf-8") as f:
        data = json.load(f)
    meta = data["_meta"]

    doc = Document()

    # Margens
    for section in doc.sections:
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)

    # ── Capa ────────────────────────────────────────────────────────
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Panorama Fiscal do Distrito Federal")
    r.font.size = Pt(22)
    r.font.bold = True
    r.font.color.rgb = COLOR_HEADER_DARK

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("Relatório Executivo")
    r.font.size = Pt(14)
    r.font.color.rgb = COLOR_HEADER_MID
    r.font.italic = True

    today = datetime.date.today().strftime("%d/%m/%Y")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"Atualização: {today}")
    r.font.size = Pt(10)
    r.font.color.rgb = COLOR_MUTED

    doc.add_paragraph()

    # ── Sumário Executivo ──────────────────────────────────────────
    add_heading(doc, "Sumário Executivo", level=1)
    rreo_closed = meta["raw_siconfi"]["rreo_closed"]
    rgf_dtp = meta["raw_siconfi"]["rgf_dtp"]
    rgf_caixa = meta["raw_siconfi"]["rgf_caixa"]
    fcdf = meta["fcdf_dotacao_atualizada"] or 0
    orc_total = meta["raw_siconfi"]["rreo_current"]["orcamento_total_dotacao"] or 0

    src_rreo_current = cite_rreo(meta["rreo_current_period"])
    src_rreo_closed = cite_rreo(meta["rreo_closed_period"])
    src_rgf = cite_rgf(meta["rgf_period"])

    add_paragraph(doc,
        f"O orçamento do Distrito Federal em 2026 é de {fmt_bi(orc_total)} (dotação atualizada própria, "
        f"{src_rreo_current}) acrescido de {fmt_bi(fcdf)} do Fundo Constitucional (FCDF, dotação "
        f"atualizada 2026 segundo o Portal da Transparência da União), totalizando "
        f"{fmt_bi(orc_total + fcdf)} para fazer frente às despesas do exercício. "
        f"Apesar do volume de recursos, o DF apresenta um quadro fiscal preocupante."
    )

    add_paragraph(doc,
        f"A nota CAPAG consolidada do DF é “C”, classificação que impede a contratação de "
        f"operações de crédito com garantia da União. O pior indicador é o de Poupança Corrente "
        f"(95,27% — limite para “B” é 95%), refletindo que praticamente toda a receita corrente "
        f"é consumida pela despesa corrente, sem margem para investir. A disponibilidade de "
        f"caixa líquida (após inscrição em Restos a Pagar) caiu de "
        f"{fmt_bi(data['caixa'][0]['caixa_liquido'])} em 2021 para "
        f"{fmt_mi(rgf_caixa['caixa_liquido_total'])} em 2025 ({src_rgf}, Anexo 05), "
        f"com recursos não vinculados negativos em "
        f"{fmt_mi(rgf_caixa['caixa_liquido_nao_vinculado'])}."
    )

    add_paragraph(doc,
        f"A taxa de investimento sobre a RCL é de 4,7%, posicionando o DF em 20º lugar "
        f"entre as 27 unidades da federação. Para alcançar a média nacional (8,4%) seriam "
        f"necessários R$ 1,5 bi/ano adicionais; para o nível do líder (Piauí, 17,0%), R$ 4,8 bi/ano. "
        f"A despesa com pessoal está em {fmt_pct(rgf_dtp['pct_pessoal_rcl_ajustada'])} da RCL "
        f"ajustada ({src_rgf}, Anexo 01) — dentro do limite legal, mas próxima do alerta (44,1%). "
        f"Os benefícios tributários totalizam R$ 13,4 bi (dado da SEEC-DF), equivalente a 57% da "
        f"receita realizada de impostos no ano ({src_rreo_closed}, Anexo 01)."
    )

    add_paragraph(doc,
        "A combinação desses sinais aponta para a necessidade urgente de uma agenda de "
        "recuperação fiscal sustentada em três eixos: (i) revisão do gasto direto; "
        "(ii) governança dos incentivos tributários; e (iii) reconstrução da poupança "
        "corrente e fortalecimento da governança fiscal.",
        bold=False
    )

    doc.add_paragraph()

    # ── Indicadores ─────────────────────────────────────────────────
    add_heading(doc, "Indicadores em destaque", level=1)
    add_kpi_table(doc, data["kpis"])
    add_source(doc,
        f"Orçamento DF: {src_rreo_current}, Anexo 01 (Balanço Orçamentário), linha "
        f"\"TOTAL DAS DESPESAS (XII) = (X + XI)\", coluna DOTAÇÃO ATUALIZADA. "
        f"FCDF: Portal da Transparência da União — órgão 25915. "
        f"Investimento/RCL: cálculo a partir de {src_rreo_closed} (Anexo 01, linha INVESTIMENTOS) "
        f"e {src_rgf} (Anexo 01, RCL). CAPAG: Tesouro Transparente. "
        f"Caixa: {src_rgf}, Anexo 05. Benefícios: Beneficiômetro SEEC-DF."
    )
    doc.add_paragraph()

    # ── Investimento ────────────────────────────────────────────────
    add_heading(doc, "Investimento Público", level=1)
    add_paragraph(doc,
        f"A taxa de investimento sobre a RCL do DF (4,7% em 2025) o coloca em 20º lugar "
        f"no ranking nacional, atrás da média BR (8,4%) e muito distante de estados como "
        f"Piauí (17%), Espírito Santo (14,8%) e Maranhão (14%). Em termos absolutos, o "
        f"investimento liquidado evoluiu de R$ 0,7 bi em 2021 para "
        f"{fmt_bi(rreo_closed['investimento_liquidado'])} em 2025."
    )
    add_source(doc,
        f"Investimento liquidado: {src_rreo_closed}, Anexo 01, linha INVESTIMENTOS, "
        f"coluna DESPESAS LIQUIDADAS ATÉ O BIMESTRE. "
        f"RCL: {src_rgf}, Anexo 01, linha RECEITA CORRENTE LÍQUIDA (IV). "
        f"Ranking entre UFs: SICONFI (compilação manual a partir do RREO do exercício)."
    )
    add_paragraph(doc, "Lacunas (gaps) para fechar:", bold=True)
    add_bullet(doc, "Para atingir a média nacional (8,4% da RCL): +R$ 1,5 bi/ano")
    add_bullet(doc, "Para atingir o nível do líder Piauí (17% da RCL): +R$ 4,8 bi/ano")
    add_bullet(doc, "Em termos per capita, o DF ocupa o 12º lugar no país")
    doc.add_paragraph()

    # ── CAPAG ──────────────────────────────────────────────────────
    add_heading(doc, "CAPAG — Capacidade de Pagamento", level=1)
    add_paragraph(doc,
        f"A nota CAPAG consolidada do DF é “{meta.get('capag_nota_consolidada','C')}”, "
        f"que restringe novas operações de crédito com garantia da União. "
        f"A CAPAG é composta por três indicadores avaliados pelo Tesouro Nacional:"
    )
    doc.add_paragraph()
    add_capag_table(doc, data["capag"], data["capag_historico"], meta.get("capag_nota_consolidada"))
    add_source(doc, "Tesouro Nacional / Tesouro Transparente — CAPAG dos Estados (publicação anual).")

    doc.add_paragraph()
    add_paragraph(doc, "Histórico (2018–2025):", bold=True)
    add_capag_hist_table(doc, data["capag_historico"])
    add_source(doc, "Tesouro Nacional / Tesouro Transparente — CAPAG dos Estados (publicações anuais 2018–2025).")
    doc.add_paragraph()
    add_paragraph(doc,
        "O DF estava na nota B (2021–2023) e voltou à C em 2024 após deterioração da "
        "poupança corrente. O indicador de Endividamento é o ponto forte: com 30,85% da "
        "RCL (limite para A é 60%), há espaço teórico de mais de R$ 11 bi em dívida — "
        "que, porém, não pode ser contratada com garantia da União enquanto a "
        "classificação consolidada não melhorar."
    )
    doc.add_paragraph()

    # ── Caixa ──────────────────────────────────────────────────────
    add_heading(doc, "Disponibilidade de Caixa", level=1)
    add_paragraph(doc,
        "A disponibilidade de caixa líquida (após inscrição em Restos a Pagar não "
        "processados) é o indicador que melhor sintetiza a folga financeira de "
        "curto prazo. A evolução é preocupante:"
    )
    doc.add_paragraph()
    add_caixa_table(doc, data["caixa"])
    add_source(doc,
        f"RGF – Poder Executivo – 3º quadrimestre de cada exercício (2021 a 2025), "
        f"Anexo 05 — Demonstrativo da Disponibilidade de Caixa e dos Restos a Pagar. "
        f"Linha TOTAL ({chr(0x201C)}IV = I+II+III{chr(0x201D)} desde 2023, {chr(0x201C)}III = I+II{chr(0x201D)} em 2021-22) e linha "
        f"TOTAL DOS RECURSOS NÃO VINCULADOS (I), coluna DISPONIBILIDADE DE CAIXA LÍQUIDA "
        f"(APÓS A INSCRIÇÃO EM RESTOS A PAGAR NÃO PROCESSADOS DO EXERCÍCIO)."
    )
    doc.add_paragraph()
    add_paragraph(doc,
        f"O caixa líquido caiu de {fmt_bi(data['caixa'][0]['caixa_liquido'])} em 2021 "
        f"para {fmt_mi(data['caixa'][-1]['caixa_liquido'])} em 2025 — queda de cerca de 70%. "
        f"Mais grave: o saldo de recursos não vinculados (a parte discricionária do caixa) "
        f"está negativo em {fmt_mi(data['caixa'][-1]['nao_vinculado'])}, o que significa "
        f"que o DF financia despesas discricionárias com recursos vinculados — "
        f"prática que comprime ainda mais a capacidade de resposta a choques."
    )
    doc.add_paragraph()

    # ── Pessoal ────────────────────────────────────────────────────
    add_heading(doc, "Despesa com Pessoal e LRF", level=1)
    add_paragraph(doc,
        f"A despesa total com pessoal do Poder Executivo está em "
        f"{fmt_pct(rgf_dtp['pct_pessoal_rcl_ajustada'])} da RCL ajustada — "
        f"abaixo dos limites de alerta (44,1%) e prudencial (46,55%), mas com "
        f"margem estreita. Os patamares legais e o espaço fiscal disponível são:"
    )
    doc.add_paragraph()
    add_pessoal_table(doc, data["pessoal"])
    add_source(doc,
        f"{src_rgf}, Anexo 01 — Demonstrativo da Despesa com Pessoal. "
        f"Linha DESPESA TOTAL COM PESSOAL – DTP (VI), coluna % SOBRE A RCL AJUSTADA."
    )
    doc.add_paragraph()
    add_paragraph(doc,
        "Embora o DF esteja dentro dos limites da LRF, qualquer movimento de "
        "ampliação relevante da folha (reajustes, novos concursos) precisa ser "
        "examinado à luz dessa margem cada vez mais estreita.",
        italic=True, color=COLOR_MUTED
    )
    doc.add_paragraph()

    # ── Benefícios Tributários ─────────────────────────────────────
    add_heading(doc, "Benefícios Tributários", level=1)
    beneficios_bi = 13.4  # do manual
    receita_imp = rreo_closed["receita_impostos_realizada"]
    pct = beneficios_bi * 1e9 / receita_imp * 100
    add_paragraph(doc,
        f"O valor total de benefícios tributários concedidos pelo DF em 2025 foi "
        f"de R$ {beneficios_bi:.1f} bi".replace(".", ",") +
        f" (Beneficiômetro SEEC-DF), o que equivale a {pct:.0f}% da receita realizada "
        f"de impostos no exercício ({fmt_bi(receita_imp)} — {src_rreo_closed}, Anexo 01, "
        f"linha Impostos, coluna RECEITAS REALIZADAS ATÉ O BIMESTRE). É um volume expressivo, "
        f"com implicações fiscais diretas: cada ponto percentual a menos de benefício pode "
        f"liberar centenas de milhões de reais para outras finalidades (caixa, investimento, pessoal)."
    )
    add_paragraph(doc,
        "Os principais campos de atenção:"
    )
    add_bullet(doc, "ICMS representa ~85% das renúncias — foco para revisão prioritária")
    add_bullet(doc, "Ausência de governança formal sobre prazos, contrapartidas e metas de cada benefício")
    add_bullet(doc, "Necessidade de avaliação periódica de custo-benefício por desoneração")
    doc.add_paragraph()

    # ── PPPs ───────────────────────────────────────────────────────
    add_heading(doc, "Parcerias Público-Privadas", level=1)
    proj = data["ppps_projecao"]
    atual_proj = next((r for r in proj if r["ano"] == "2025"), proj[0])
    pico_proj = max(proj, key=lambda r: r["pct"])
    add_paragraph(doc,
        f"O comprometimento da RCL com PPPs em 2025 é de "
        f"{fmt_pct(atual_proj['pct'], 2)} ({fmt_mi(atual_proj['despesas_ppp'])}). "
        f"A projeção indica um pico de {fmt_pct(pico_proj['pct'], 2)} em "
        f"{pico_proj['ano']} (R$ {pico_proj['despesas_ppp']/1e9:.2f} bi), "
        f"ainda dentro do limite legal de 5% da RCL (Lei 11.079/2004) — mas com pouca margem "
        f"se novas PPPs forem contratadas sem ajustes compensatórios."
    )
    doc.add_paragraph()
    add_paragraph(doc, "Carteira de projetos:", bold=True)
    add_ppp_table(doc, data["ppps"])
    add_source(doc,
        f"Lista de projetos e projeção de comprometimento: Demonstrativo das Parcerias "
        f"Público-Privadas do {src_rreo_closed}, Anexo 13."
    )
    doc.add_paragraph()

    # ── Agenda ─────────────────────────────────────────────────────
    add_heading(doc, "Agenda de Recuperação Fiscal", level=1)
    add_paragraph(doc,
        "Os três eixos prioritários para a recuperação fiscal do DF são:"
    )
    # group agenda by eixo
    eixos_order = []
    eixos = {}
    for row in data["agenda"]:
        e = row["eixo"]
        if e not in eixos:
            eixos[e] = {"titulo": row["titulo"], "itens": []}
            eixos_order.append(e)
        eixos[e]["itens"].append(row["item"])

    for e in eixos_order:
        info = eixos[e]
        h = doc.add_heading(f"Eixo {e}. {info['titulo']}", level=2)
        for r in h.runs:
            r.font.color.rgb = COLOR_HEADER_MID
            r.font.size = Pt(12)
        for it in info["itens"]:
            add_bullet(doc, it)
    doc.add_paragraph()

    # ── Considerações finais ───────────────────────────────────────
    add_heading(doc, "Considerações Finais", level=1)
    add_paragraph(doc,
        "O DF possui ativos institucionais relevantes: orçamento robusto somado ao FCDF, "
        "endividamento controlado, e um conjunto razoável de PPPs estruturadas e em "
        "estruturação. O ponto crítico está na poupança corrente: enquanto a despesa "
        "corrente consumir 95%+ da receita corrente, não haverá folga para investimento, "
        "nem capacidade de absorver choques, nem retomada da nota CAPAG."
    )
    add_paragraph(doc,
        "A combinação das três frentes da Agenda — revisão do gasto direto, governança "
        "dos benefícios tributários e reconstrução da poupança — é, portanto, o caminho "
        "estrutural para recolocar o DF em situação fiscal saudável. Ações pontuais "
        "(ajuste de curto prazo, contingenciamento) podem mitigar pressões imediatas mas "
        "não substituem o ajuste estrutural necessário."
    )
    doc.add_paragraph()

    # ── Fontes ─────────────────────────────────────────────────────
    add_heading(doc, "Fontes e Metodologia", level=1)
    add_paragraph(doc,
        f"Atualização dos dados: {meta['updated_at']}", italic=True, color=COLOR_MUTED
    )
    add_bullet(doc, f"SICONFI/STN — {src_rreo_current} (orçamento DF 2026)")
    add_bullet(doc, f"SICONFI/STN — {src_rreo_closed} (investimento liquidado, receita de impostos, PPPs)")
    add_bullet(doc, f"SICONFI/STN — {src_rgf} (RCL, DTP e disponibilidade de caixa)")
    add_bullet(doc, "SICONFI/STN — RGF – Poder Executivo – 3º quadrimestre dos exercícios 2021 a 2024 (histórico de caixa)")
    add_bullet(doc, "Tesouro Transparente — CAPAG dos Estados (publicações anuais 2018–2025)")
    add_bullet(doc, "IBGE — população residente estimada (SIDRA, agregado 6579)")
    add_bullet(doc, "Portal da Transparência da União — dotação atualizada do FCDF (órgão 25915)")
    add_bullet(doc, "Beneficiômetro SEEC-DF — valor total dos benefícios tributários")
    add_paragraph(doc,
        "Painel ao vivo: https://couri-consult.github.io/panorama-fiscal-df/",
        italic=True, color=COLOR_MUTED
    )

    # save
    os.makedirs(os.path.dirname(OUTPUT), exist_ok=True)
    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT} ({os.path.getsize(OUTPUT):,} bytes)")


if __name__ == "__main__":
    build_report()
